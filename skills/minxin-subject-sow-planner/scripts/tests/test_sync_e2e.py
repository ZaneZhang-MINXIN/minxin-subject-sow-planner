#!/usr/bin/env python3

from __future__ import annotations

import sys
import tempfile
import unittest
import os
import zipfile
from pathlib import Path

from docx import Document

ROOT = Path(__file__).resolve().parents[2]
sys.path.insert(0, str(ROOT / "scripts"))

from export_sow import create_sow  # noqa: E402
from planner_core import build_planner_data, read_json, write_json  # noqa: E402
from sow_planner import run_node  # noqa: E402
from validate import validate_docx  # noqa: E402


class SyncEndToEndTests(unittest.TestCase):
    def test_exported_workbook_persists_freeze_rows(self):
        node_value = os.environ.get("MINXIN_TEST_NODE")
        modules_value = os.environ.get("MINXIN_TEST_NODE_MODULES")
        if not node_value or not modules_value:
            self.skipTest("Set MINXIN_TEST_NODE and MINXIN_TEST_NODE_MODULES")
        calendar = read_json(ROOT / "assets" / "ay2026-27-calendar.json")
        fixture = read_json(ROOT / "scripts" / "tests" / "fixtures" / "multi-subject-fixture.json")
        payload = build_planner_data(fixture, calendar)
        with tempfile.TemporaryDirectory(prefix="minxin-freeze-e2e-") as name:
            directory = Path(name)
            data, workbook = directory / "planner.json", directory / "planner.xlsx"
            write_json(data, payload)
            run_node(ROOT / "scripts" / "build_planner.mjs", ["--data", str(data), "--out", str(workbook)], Path(node_value), Path(modules_value))
            roundtrip = directory / "roundtrip.json"
            run_node(ROOT / "scripts" / "read_planner.mjs", ["--input", str(workbook), "--out", str(roundtrip)], Path(node_value), Path(modules_value))
            self.assertEqual(read_json(roundtrip)["source_hash"], payload["source_hash"])
            with zipfile.ZipFile(workbook) as archive:
                sheets = [member for member in archive.namelist() if member.startswith("xl/worksheets/sheet") and member.endswith(".xml")]
                self.assertEqual(len(sheets), 10)
                self.assertTrue(all(b'ySplit="4"' in archive.read(member) and b'state="frozen"' in archive.read(member) for member in sheets))

    def test_one_objective_change_updates_only_linked_course_and_both_profiles(self):
        calendar = read_json(ROOT / "assets" / "ay2026-27-calendar.json")
        fixture = read_json(ROOT / "scripts" / "tests" / "fixtures" / "multi-subject-fixture.json")
        baseline = build_planner_data(fixture, calendar)
        marker = "SYNC-MARKER: cites two precise details and explains their combined effect."
        target = next(row for row in fixture["tables"]["Objectives"] if row["objective_id"] == "ENGLISH-G7-CORE-A-U1-O1")
        target["objective_text"] = marker
        linked_plan = next(row for row in fixture["tables"]["Weekly_Plan"] if row["plan_id"] == "PLAN-ENGLISH-G7-CORE-A-W01")
        linked_plan["major_concern_refs"] = "MC2"
        updated = build_planner_data(fixture, calendar)
        self.assertNotEqual(baseline["source_hash"], updated["source_hash"])
        english_rows = [row for row in updated["tables"]["SOW_View"] if row["course_id"] == "ENGLISH-G7-CORE-A"]
        other_rows = [row for row in updated["tables"]["SOW_View"] if row["course_id"] != "ENGLISH-G7-CORE-A"]
        self.assertTrue(any(marker in row["teaching_objectives"] for row in english_rows))
        self.assertFalse(any(marker in row["teaching_objectives"] for row in other_rows))
        course = next(row for row in updated["tables"]["Course_Brief"] if row["course_id"] == "ENGLISH-G7-CORE-A")
        with tempfile.TemporaryDirectory(prefix="minxin-sync-e2e-") as name:
            directory = Path(name)
            standard = directory / "standard.docx"
            compact = directory / "compact.docx"
            create_sow(updated, course, english_rows, standard, "standard", "en-GB", ROOT / "assets" / "minxin-logo.jpeg")
            create_sow(updated, course, english_rows, compact, "compact", "en-GB", ROOT / "assets" / "minxin-logo.jpeg")
            standard_doc, compact_doc = Document(standard), Document(compact)
            self.assertTrue(any(len(table.columns) == 11 for table in standard_doc.tables))
            self.assertTrue(any(len(table.columns) == 10 for table in compact_doc.tables))
            text = "\n".join(cell.text for table in standard_doc.tables for row in table.rows for cell in row.cells)
            self.assertIn(marker, text)
            sow_table = next(table for table in standard_doc.tables if len(table.columns) == 11)
            self.assertIn("(M2)", sow_table.rows[1].cells[5].text)
            self.assertNotRegex(sow_table.rows[2].cells[5].text, r"\(M[123]")
            outside_text = "\n".join(paragraph.text for paragraph in standard_doc.paragraphs)
            self.assertIn("Remarks: Major Concerns (2025–2028)", outside_text)
            self.assertIn("M1: To enhance global engagement and innovation", outside_text)
            self.assertIn("M2: To promote high quality teaching and academic excellence", outside_text)
            self.assertIn("M3: To foster students’ holistic development and global outlook", outside_text)
            with zipfile.ZipFile(standard) as archive:
                document_xml = archive.read("word/document.xml")
                settings_xml = archive.read("word/settings.xml")
            self.assertNotIn(b"<w:br", document_xml)
            self.assertNotIn(b"<w:cr", document_xml)
            self.assertNotIn(b"<w:noWrap", document_xml)
            self.assertNotIn(b"<w:softHyphen", document_xml)
            self.assertIn(b"<w:wordWrap", document_xml)
            self.assertIn(b"<w:suppressAutoHyphens", document_xml)
            self.assertIn(b'<w:autoHyphenation w:val="0"', settings_xml)
            self.assertIn("\u200b", standard_doc.tables[1].cell(1, 10).text)
            self.assertEqual(validate_docx(standard, course, english_rows, updated["source_hash"], "standard", "en-GB"), [])


if __name__ == "__main__":
    unittest.main()
