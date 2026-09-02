#!/usr/bin/env python3
"""Validate planner relationships, generated lineage, Word synchronization, and release evidence."""

from __future__ import annotations

import argparse
import json
import re
from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from PIL import Image, ImageStat

from export_sow import MAJOR_CONCERNS, COMPACT_WIDTHS, EN_HEADERS, STANDARD_WIDTHS, ZH_HEADERS, assessment_text, row_values
from planner_core import ALL_SHEETS, HEADERS, normalize_text, read_json, refresh_planner_payload, validate_tables, write_json


def compare_headers(payload: dict) -> list[str]:
    errors = []
    tables = payload.get("tables", {})
    missing = [name for name in ALL_SHEETS if name not in tables]
    if missing:
        errors.append("Missing sheets: " + ", ".join(missing))
    for name in ALL_SHEETS:
        actual = payload.get("headers", {}).get(name, [])
        if actual and actual != HEADERS[name]:
            errors.append(f"{name}: columns differ from the v2.1 schema")
    return errors


def normalize(value) -> str:
    return re.sub(r"\s+", " ", str(value or "").replace("\u200b", "")).strip()


def find_sow_table(document: Document, expected_columns: int):
    return next((table for table in document.tables if len(table.columns) == expected_columns), None)


def document_paragraphs(document: Document):
    yield from document.paragraphs
    seen_cells = set()
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                identity = id(cell._tc)
                if identity in seen_cells:
                    continue
                seen_cells.add(identity)
                yield from cell.paragraphs
    for section in document.sections:
        yield from section.footer.paragraphs


def validate_soft_wrapping(path: Path, document: Document) -> list[str]:
    errors = []
    for index, paragraph in enumerate(document_paragraphs(document), 1):
        p_pr = paragraph._p.pPr
        word_wrap = p_pr.find(qn("w:wordWrap")) if p_pr is not None else None
        suppress = p_pr.find(qn("w:suppressAutoHyphens")) if p_pr is not None else None
        if word_wrap is None or word_wrap.get(qn("w:val")) not in {"1", "true", "on"}:
            errors.append(f"{path.name}: paragraph {index} does not enable automatic word wrapping")
        if suppress is None or suppress.get(qn("w:val")) not in {"1", "true", "on"}:
            errors.append(f"{path.name}: paragraph {index} does not suppress automatic hyphenation")
    with ZipFile(path) as archive:
        document_xml = archive.read("word/document.xml")
        settings_xml = archive.read("word/settings.xml")
        for tag in (b"<w:br", b"<w:cr", b"<w:noWrap", b"<w:softHyphen"):
            if tag in document_xml:
                errors.append(f"{path.name}: prohibited Word control {tag.decode()} found in document.xml")
        if re.search(rb"<w:trHeight\b[^>]*w:hRule=\"exact\"", document_xml):
            errors.append(f"{path.name}: exact row height can clip editable text")
        if not re.search(rb"<w:autoHyphenation\b[^>]*w:val=\"(?:0|false|off)\"", settings_xml):
            errors.append(f"{path.name}: document-level automatic hyphenation is not disabled")
    return errors


def validate_major_concerns_remarks(path: Path, document: Document, sow_table) -> list[str]:
    errors = []
    outside_paragraphs = document.paragraphs
    heading = next((paragraph for paragraph in outside_paragraphs if normalize(paragraph.text) == "Remarks: Major Concerns (2025–2028)"), None)
    if heading is None:
        return [f"{path.name}: missing table-external Remarks: Major Concerns (2025–2028)"]
    if document.element.body.index(heading._p) <= document.element.body.index(sow_table._tbl):
        errors.append(f"{path.name}: Major Concerns Remarks must follow the main SOW table")
    outside_text = " ".join(normalize(paragraph.text) for paragraph in outside_paragraphs)
    for code, statement in MAJOR_CONCERNS:
        if normalize(f"{code}: {statement}") not in outside_text:
            errors.append(f"{path.name}: Remarks omit or alter the official {code} statement")
    return errors


def validate_docx(path: Path, course: dict, rows: list[dict], source_hash: str, profile: str, language: str) -> list[str]:
    errors = []
    if not path.exists():
        return [f"Missing Word output: {path.name}"]
    document = Document(path)
    expected_columns = 10 if profile == "compact" else 11
    if len(document.sections) != 1 or document.sections[0].orientation != WD_ORIENT.LANDSCAPE:
        errors.append(f"{path.name}: expected one A4 landscape section")
    table = find_sow_table(document, expected_columns)
    if table is None:
        return errors + [f"{path.name}: no {expected_columns}-column SOW table"]
    errors.extend(validate_soft_wrapping(path, document))
    errors.extend(validate_major_concerns_remarks(path, document, table))
    expected_headers = (ZH_HEADERS if language == "zh-Hant-HK" else EN_HEADERS).copy()
    if profile == "compact":
        del expected_headers[4]
    if [normalize(cell.text) for cell in table.rows[0].cells] != expected_headers:
        errors.append(f"{path.name}: header mismatch")
    if len(table.rows) - 1 != len(rows):
        return errors + [f"{path.name}: expected {len(rows)} data rows, found {len(table.rows)-1}"]
    for index, (source, word_row) in enumerate(zip(rows, table.rows[1:]), 1):
        actual = [normalize(cell.text) for cell in word_row.cells]
        expected = [normalize(value) for value in row_values(source, profile == "compact")]
        if source.get("row_type") == "CALENDAR_BLOCK":
            if actual[0] != expected[0] or actual[1] != expected[1] or actual[2] != normalize(source.get("module_unit")):
                errors.append(f"{path.name}: calendar row {index} differs from SOW_View")
            if len({id(cell._tc) for cell in word_row.cells}) != 3:
                errors.append(f"{path.name}: calendar row {index} is not Date | Week | merged label")
        elif actual != expected:
            errors.append(f"{path.name}: instructional row {index} differs from SOW_View")
    keywords = document.core_properties.keywords or ""
    if f"source_hash={source_hash}" not in keywords or f"course_id={course.get('course_id')}" not in keywords:
        errors.append(f"{path.name}: missing or stale source lineage metadata")
    return errors


def render_findings(render_dir: Path) -> list[str]:
    errors = []
    pages = sorted(render_dir.rglob("page-*.png"))
    if not pages:
        return ["No rendered Word pages found"]
    for page in pages:
        image = Image.open(page).convert("L")
        if image.width <= image.height:
            errors.append(f"{page}: rendered page is not landscape")
        mean = ImageStat.Stat(image.resize((200, 150))).mean[0]
        if mean > 253.8:
            errors.append(f"{page}: appears blank")
    return errors


def release_blocking_issues(issues: list[dict]) -> list[dict]:
    return [
        issue for issue in issues
        if issue.get("status") == "OPEN" and str(issue.get("stop_condition", "")).strip()
    ]


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--planner-json", type=Path, required=True)
    parser.add_argument("--word-dir", type=Path)
    parser.add_argument("--profile", choices=["standard", "compact"], default="standard")
    parser.add_argument("--language", choices=["en-GB", "zh-Hant-HK"])
    parser.add_argument("--report", type=Path)
    parser.add_argument("--release", action="store_true")
    parser.add_argument("--render-dir", type=Path)
    args = parser.parse_args()
    payload = read_json(args.planner_json)
    errors = compare_headers(payload)
    refreshed = refresh_planner_payload(payload)
    if str(payload.get("source_hash", "")) != refreshed["source_hash"]:
        errors.append("STALE_SOURCE_HASH: top-level planner source hash does not match editable tables")
    expected_view = refreshed["tables"]["SOW_View"]
    current_issues = validate_tables(refreshed["tables"])
    stale_issues = validate_tables(
        payload["tables"], expected_view=expected_view,
        expected_weeks=refreshed["tables"]["Weeks"],
        expected_qa=refreshed["tables"]["QA"],
        expected_source_hash=refreshed["source_hash"],
    )
    issues = list({issue["qa_id"]: issue for issue in [*current_issues, *stale_issues]}.values())
    high = [issue for issue in issues if issue["severity"] == "HIGH"]
    errors.extend(f"{issue['code']}: {issue['message']} [{issue['record_refs']}]" for issue in high)
    release_blockers = release_blocking_issues(issues)
    word_errors = []
    if args.word_dir:
        for course in refreshed["tables"]["Course_Brief"]:
            language = args.language or course.get("output_language") or "en-GB"
            suffix = "ZH-HANT-HK" if language == "zh-Hant-HK" else "EN-GB"
            safe = re.sub(r"[^A-Za-z0-9_-]+", "_", str(course["course_id"])).strip("_")
            path = args.word_dir / f"MINXIN_{safe}_SOW_{suffix}.docx"
            rows = [row for row in expected_view if row.get("course_id") == course.get("course_id")]
            word_errors.extend(validate_docx(path, course, rows, refreshed["source_hash"], args.profile, language))
    errors.extend(word_errors)
    if args.release:
        errors.extend(
            f"OPEN_RELEASE_QA {issue['code']}: {issue['message']} [{issue['record_refs']}]"
            for issue in release_blockers if issue["severity"] != "HIGH"
        )
        if not args.word_dir or not args.render_dir:
            errors.append("Release validation requires --word-dir and --render-dir")
        else:
            errors.extend(render_findings(args.render_dir))
    report = {
        "status": "PASS" if not errors else "FAIL",
        "release_ready": bool(args.release and not errors),
        "source_hash": refreshed["source_hash"],
        "error_count": len(errors),
        "errors": errors,
        "qa_counts": {
            "HIGH": sum(1 for issue in issues if issue["severity"] == "HIGH"),
            "MEDIUM": sum(1 for issue in issues if issue["severity"] == "MEDIUM"),
        },
        "open_release_qa_count": len(release_blockers),
        "open_release_qa_codes": sorted({issue["code"] for issue in release_blockers}),
        "word_files_checked": len(list(args.word_dir.glob("*.docx"))) if args.word_dir and args.word_dir.exists() else 0,
    }
    if args.report:
        write_json(args.report, report)
    print(json.dumps(report, ensure_ascii=False))
    return 0 if not errors else 1


if __name__ == "__main__":
    raise SystemExit(main())
