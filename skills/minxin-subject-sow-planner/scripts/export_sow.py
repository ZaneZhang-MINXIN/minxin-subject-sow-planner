#!/usr/bin/env python3
"""Export one MINXIN Word SOW per course from a refreshed planner JSON."""

from __future__ import annotations

import argparse
import re
import tempfile
from pathlib import Path
from zipfile import ZIP_DEFLATED, ZipFile

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Mm, Pt, RGBColor

from planner_core import read_json


EN_HEADERS = [
    "Date", "Week", "Module / Unit", "Learning Focus", "Prior Learning", "Learning Objectives",
    "Periods", "Resources", "Values / National Education", "Learning & Teaching Activities",
    "Assessment / Evidence & Feedback",
]
ZH_HEADERS = [
    "日期", "學周", "單元", "學習重點", "已有知識", "學習目標", "課節", "資源",
    "價值觀／國民教育", "學與教活動", "評估／證據與回饋",
]
STANDARD_WIDTHS = [0.72, 0.42, 0.72, 0.75, 0.68, 1.20, 0.40, 0.65, 0.62, 1.65, 1.88]
COMPACT_WIDTHS = [0.76, 0.44, 0.78, 0.82, 1.35, 0.42, 0.72, 0.68, 1.75, 1.97]


def scrub_revision_session_ids(path: Path) -> None:
    """Remove Word editing-session identifiers without changing visible content."""
    with tempfile.NamedTemporaryFile(
        prefix=f".{path.stem}-", suffix=".docx", dir=path.parent, delete=False
    ) as handle:
        scrubbed_path = Path(handle.name)
    try:
        with ZipFile(path) as source, ZipFile(scrubbed_path, "w", ZIP_DEFLATED) as target:
            for item in source.infolist():
                data = source.read(item.filename)
                if item.filename.startswith("word/") and item.filename.endswith(".xml"):
                    data = re.sub(rb'\s+w:rsid[A-Za-z]*="[^"]*"', b"", data)
                    data = re.sub(rb"<w:rsids\b[^>]*>.*?</w:rsids>", b"", data, flags=re.DOTALL)
                target.writestr(item, data)
        scrubbed_path.replace(path)
    finally:
        if scrubbed_path.exists():
            scrubbed_path.unlink()


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_margins(cell, top: int = 45, start: int = 55, bottom: int = 45, end: int = 55) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    element = OxmlElement("w:tblHeader")
    element.set(qn("w:val"), "true")
    tr_pr.append(element)


def set_table_widths(table, widths: list[float]) -> None:
    """Persist fixed grid widths so Word and LibreOffice render the same layout."""
    for index, width in enumerate(widths):
        table.columns[index].width = Inches(width)
    grid_columns = table._tbl.tblGrid.gridCol_lst
    for grid_column, width in zip(grid_columns, widths):
        grid_column.w = Inches(width)


def prevent_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tr_pr.append(OxmlElement("w:cantSplit"))


def remove_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = OxmlElement(f"w:{edge}")
        tag.set(qn("w:val"), "nil")
        borders.append(tag)


def set_font(run, size: float, bold: bool = False, colour: str = "000000") -> None:
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(colour)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    set_font(run, 9)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, end])


def configure_document(document: Document, title: str, course_id: str, source_hash: str) -> None:
    section = document.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = Mm(297), Mm(210)
    section.top_margin = Mm(12.5)
    section.bottom_margin = Mm(12.5)
    section.left_margin = Mm(25.4)
    section.right_margin = Mm(25.4)
    section.header_distance = Mm(5)
    section.footer_distance = Mm(6)
    styles = document.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
    styles["Normal"].font.size = Pt(8)
    props = document.core_properties
    props.title = title
    props.subject = "Scheme of Work"
    props.author = "MINXIN School"
    props.last_modified_by = "MINXIN School"
    props.keywords = f"course_id={course_id};source_hash={source_hash}"
    props.comments = "Generated from the MINXIN subject SOW planner single source of truth."
    footer = section.footer
    add_page_number(footer.paragraphs[0])


def add_title(document: Document, title: str, subtitle: str, logo: Path | None) -> None:
    table = document.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(1.05)
    table.columns[1].width = Inches(8.6)
    remove_table_borders(table)
    left, right = table.rows[0].cells
    if logo and logo.exists():
        paragraph = left.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.add_run().add_picture(str(logo), width=Inches(0.72))
    paragraph = right.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(title)
    set_font(run, 16, True)
    paragraph.paragraph_format.space_after = Pt(1)
    second = document.add_paragraph()
    second.alignment = WD_ALIGN_PARAGRAPH.CENTER
    second.paragraph_format.space_before = Pt(0)
    second.paragraph_format.space_after = Pt(5)
    run = second.add_run(subtitle)
    set_font(run, 10.5, True)


def assessment_text(row: dict) -> str:
    parts = []
    if row.get("evidence"):
        parts.append(f"Evidence: {row['evidence']}")
    if row.get("assessment_purpose"):
        parts.append(str(row["assessment_purpose"]))
    if row.get("feedback_revision"):
        parts.append(f"Feedback/revision: {row['feedback_revision']}")
    return "\n".join(parts)


def row_values(row: dict, compact: bool) -> list[str]:
    values = [
        row.get("date_range", ""), row.get("week_label", ""), row.get("module_unit", ""),
        row.get("learning_focus", ""), row.get("prior_learning", ""), row.get("teaching_objectives", ""),
        row.get("periods", ""), row.get("resources", ""), row.get("values", ""),
        row.get("activities", ""), assessment_text(row),
    ]
    if compact:
        del values[4]
    return [str(value) for value in values]


def create_sow(planner: dict, course: dict, rows: list[dict], output: Path, profile: str, language: str, logo: Path | None) -> None:
    compact = profile == "compact"
    headers = (ZH_HEADERS if language == "zh-Hant-HK" else EN_HEADERS).copy()
    widths = COMPACT_WIDTHS if compact else STANDARD_WIDTHS
    if compact:
        del headers[4]
    academic_year = next((str(row.get("value")) for row in planner["tables"]["Setup"] if row.get("key") == "academic_year"), "")
    grade = str(course.get("grade_level", "")).strip()
    title = f"{grade + ' ' if grade else ''}{course.get('subject_name', 'SUBJECT')} SCHEME OF WORK {academic_year}".upper()
    subtitle = f"Course: {course.get('course_id')} | Level/Class: {course.get('course_level') or '—'} / {course.get('class_id') or '—'}"
    document = Document()
    configure_document(document, title, str(course.get("course_id")), str(planner.get("source_hash", "")))
    add_title(document, title, subtitle, logo)
    table = document.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.style = "Table Grid"
    set_table_widths(table, widths)
    for index, (cell, header, width) in enumerate(zip(table.rows[0].cells, headers, widths)):
        cell.width = Inches(width)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(cell, "FFFF99")
        set_cell_margins(cell, 35, 20, 35, 20)
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(header)
        set_font(run, 7.6, True, "0000FF")
    table.rows[0].height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
    set_repeat_header(table.rows[0])
    for source in rows:
        cells = table.add_row().cells
        prevent_split(table.rows[-1])
        if source.get("row_type") == "CALENDAR_BLOCK":
            values = row_values(source, compact)
            for index in (0, 1):
                cells[index].text = values[index]
            merged = cells[2]
            for cell in cells[3:]:
                merged = merged.merge(cell)
            merged.text = source.get("module_unit", "Calendar event")
            for cell in table.rows[-1].cells:
                set_cell_shading(cell, "E7E6E6")
                set_cell_margins(cell, 70, 60, 70, 60)
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    paragraph.paragraph_format.space_after = Pt(0)
                    for run in paragraph.runs:
                        set_font(run, 8, True)
            continue
        values = row_values(source, compact)
        for index, (cell, text) in enumerate(zip(cells, values)):
            cell.width = Inches(widths[index])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            set_cell_margins(cell)
            cell.text = text
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 0.92
                if index in {0, 1, 6 if not compact else 5}:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    set_font(run, 7.2)
    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)
    scrub_revision_session_ids(output)


def safe_name(value: str) -> str:
    return re.sub(r"[^A-Za-z0-9_-]+", "_", value).strip("_")


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--planner-json", type=Path, required=True)
    parser.add_argument("--out-dir", type=Path, required=True)
    parser.add_argument("--course", default="all")
    parser.add_argument("--profile", choices=["standard", "compact"], default="standard")
    parser.add_argument("--language", choices=["en-GB", "zh-Hant-HK"])
    parser.add_argument("--logo", type=Path)
    args = parser.parse_args()
    planner = read_json(args.planner_json)
    courses = planner["tables"]["Course_Brief"]
    selected = courses if args.course == "all" else [course for course in courses if course.get("course_id") == args.course]
    if not selected:
        raise SystemExit(f"Unknown course: {args.course}")
    for course in selected:
        rows = [row for row in planner["tables"]["SOW_View"] if row.get("course_id") == course.get("course_id")]
        language = args.language or course.get("output_language") or "en-GB"
        suffix = "ZH-HANT-HK" if language == "zh-Hant-HK" else "EN-GB"
        output = args.out_dir / f"MINXIN_{safe_name(str(course['course_id']))}_SOW_{suffix}.docx"
        create_sow(planner, course, rows, output, args.profile, language, args.logo)
        print(output)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
