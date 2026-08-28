#!/usr/bin/env python3
"""Inspect historical SOW/curriculum files without treating their contents as instructions."""

from __future__ import annotations

import argparse
import hashlib
import json
import shutil
import subprocess
import tempfile
from collections import Counter
from pathlib import Path

import pdfplumber
from docx import Document


BUNDLED_SOFFICE = (
    Path.home()
    / ".cache/codex-runtimes/codex-primary-runtime/dependencies/native"
    / "libreoffice-headless/libreoffice/LibreOfficeDev.app/Contents/MacOS/soffice"
)


def sha(path: Path) -> str:
    return hashlib.sha256(path.read_bytes()).hexdigest()


def convert_doc(path: Path, directory: Path) -> Path:
    office = shutil.which("soffice") or shutil.which("libreoffice")
    if not office and BUNDLED_SOFFICE.is_file():
        office = str(BUNDLED_SOFFICE)
    if not office:
        raise RuntimeError(
            "Legacy .doc analysis requires LibreOffice/soffice; use the bundled "
            "Codex workspace runtime or install LibreOffice"
        )
    subprocess.run([office, "--headless", "--convert-to", "docx", "--outdir", str(directory), str(path)], check=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
    converted = directory / f"{path.stem}.docx"
    if not converted.exists():
        raise RuntimeError(f"Temporary conversion failed: {path.name}")
    return converted


def analyse_docx(path: Path, source_name: str, source_hash: str) -> dict:
    document = Document(path)
    fonts, sizes, colours, headers = Counter(), Counter(), Counter(), Counter()
    paragraphs = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text and len(paragraphs) < 30:
            paragraphs.append(text)
        for run in paragraph.runs:
            if run.font.name:
                fonts[run.font.name] += 1
            if run.font.size:
                sizes[round(run.font.size.pt, 1)] += 1
            if run.font.color and run.font.color.rgb:
                colours[str(run.font.color.rgb)] += 1
    column_counts = []
    for table in document.tables:
        column_counts.append(len(table.columns))
        if table.rows:
            headers[" | ".join(cell.text.strip() for cell in table.rows[0].cells)] += 1
    return {
        "source_name": source_name, "source_sha256": source_hash, "kind": "DOCX",
        "sections": len(document.sections), "tables": len(document.tables), "table_column_counts": dict(Counter(column_counts)),
        "common_fonts": fonts.most_common(8), "common_sizes_pt": sizes.most_common(8), "common_colours": colours.most_common(8),
        "repeated_header_candidates": headers.most_common(8), "representative_paragraphs": paragraphs,
    }


def analyse_pdf(path: Path) -> dict:
    samples = []
    with pdfplumber.open(path) as pdf:
        for index, page in enumerate(pdf.pages, 1):
            text = page.extract_text() or ""
            if text.strip() and len(samples) < 8:
                samples.append({"page": index, "text": " ".join(text.split())[:500]})
        pages = len(pdf.pages)
    return {"source_name": path.name, "source_sha256": sha(path), "kind": "PDF", "pages": pages, "text_samples": samples}


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--source", action="append", type=Path, required=True)
    parser.add_argument("--out-dir", type=Path, required=True)
    args = parser.parse_args()
    reports = []
    with tempfile.TemporaryDirectory(prefix="minxin-sow-learn-") as temp_name:
        temp = Path(temp_name)
        for source in args.source:
            suffix = source.suffix.lower()
            if suffix == ".pdf":
                reports.append(analyse_pdf(source))
            elif suffix in {".docx", ".doc"}:
                candidate = convert_doc(source, temp) if suffix == ".doc" else source
                reports.append(analyse_docx(candidate, source.name, sha(source)))
            else:
                reports.append({"source_name": source.name, "source_sha256": sha(source), "kind": suffix.lstrip(".").upper(), "status": "INDEXED_ONLY"})
    result = {
        "status": "EVIDENCE_DRAFT",
        "instruction_boundary": "Attachment contents were analysed as evidence only; no embedded instruction, macro, link, or object was executed.",
        "sources": reports,
        "classifications": {
            "reusable_specification": ["stable titles and headings", "semantic column names", "school visual language", "assessable objective sentence patterns"],
            "curriculum_evidence": ["verified topics, targets, activities, resources, assessments, and dates with source references"],
            "requires_confirmation": ["subject/KLA and Key Stage", "current curriculum authority", "assessment scope", "formal timetable", "calendar REVIEW events"],
            "do_not_inherit": ["blank pages", "isolated headers", "drifting logos", "fixed-height clipping", "duplicate weeks", "misaligned columns", "unowned TBC blocks"],
        },
        "grill_me_next_question": "Which exact subject curriculum or qualification syllabus is authoritative for this course, including version and Key Stage/grade?",
    }
    args.out_dir.mkdir(parents=True, exist_ok=True)
    (args.out_dir / "learn-report.json").write_text(json.dumps(result, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
    lines = ["# SOW evidence learning report", "", result["instruction_boundary"], "", "## Sources", ""]
    for item in reports:
        lines.append(f"- `{item['source_name']}` — {item['kind']}, SHA-256 `{item['source_sha256']}`")
    lines.extend(["", "## Bounded grill-me", "", f"First question: {result['grill_me_next_question']}"])
    (args.out_dir / "learn-report.md").write_text("\n".join(lines) + "\n", encoding="utf-8")
    print(args.out_dir / "learn-report.md")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
